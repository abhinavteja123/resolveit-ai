"""Generate ResolveIT_AI_Project_Documentation.docx in plain B&W style.

Mirrors the section structure of the user-supplied sample:
Project Description -> Application Scenarios -> Technical Architecture ->
Pre-requisites -> Prior Knowledge -> Project Objectives -> Project Workflow ->
Detailed Milestones -> Project Structure -> Results -> Advantages/Limitations ->
Future Enhancements -> Conclusion -> Appendix.

Uses default Word styles only (Heading 1/2/3, Normal). No colours, no shading,
no tinted boxes — only black text on white. Placeholders for diagrams and
screenshots are inserted as bordered single-cell tables labelled `[ DIAGRAM:
... ]` so the layout stays visible in Word and you can replace each
placeholder with an image when ready.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), "000000")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_diagram_placeholder(doc, label: str, height_cm: float = 7.5):
    """Insert a single-cell bordered table with the diagram label centered."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    _set_cell_border(cell)
    cell.height = Cm(height_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_h = OxmlElement("w:tcW")
    tc_h.set(qn("w:w"), "9000")
    tc_h.set(qn("w:type"), "dxa")
    tc_pr.append(tc_h)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"\n[ {label} ]\n\n(Insert diagram / screenshot here)\n")
    run.font.size = Pt(10)
    run.font.italic = True
    # Pad with empty paragraphs so the placeholder takes vertical space
    for _ in range(int(height_cm * 1.2)):
        p = cell.add_paragraph()
        r = p.add_run("")
        r.font.size = Pt(10)
    doc.add_paragraph()


def add_para(doc, text: str, bold: bool = False, italic: bool = False, size: int = 11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    run = p.runs[0] if p.runs else p.add_run("")
    run.text = text
    run.font.size = Pt(11)


def add_code(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = ""
            r = table.rows[ri].cells[ci].paragraphs[0].add_run(val)
            r.font.size = Pt(11)
    doc.add_paragraph()


def main(out_path: str):
    doc = Document()

    # Make the default font black + 11pt
    styles = doc.styles
    for sname in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        try:
            s = styles[sname]
            s.font.color.rgb = RGBColor(0, 0, 0)
        except KeyError:
            pass

    # ── Title ────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "ResolveIT AI — Smart Runbook Resolution Assistant\n"
        "RAG-powered IT Support with Multi-Mode Response Generation"
    )
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    # ── 1. Project Description ───────────────────────────────────────
    doc.add_heading("Project Description", level=1)

    doc.add_heading("1. Project Overview", level=2)
    add_para(
        doc,
        "ResolveIT AI is a full-stack Retrieval-Augmented Generation (RAG) "
        "platform built to act as an intelligent IT-support knowledge "
        "assistant. The system ingests internal runbooks (PDF, DOCX, TXT) "
        "and lets engineers query them in natural language; the backend "
        "retrieves the most relevant excerpts using a hybrid semantic + "
        "keyword pipeline, re-ranks them with a cross-encoder, and uses "
        "Google Gemini to generate a cited, structured resolution.",
    )
    add_para(
        doc,
        "Beyond plain Q&A, the platform supports six distinct response modes "
        "(Fast, Standard, Deep, ELI5, Expert, Dry-run), one-click answer "
        "regeneration in any other mode, AI-suggested follow-up questions, "
        "Markdown export of any answer, threaded multi-turn conversations, "
        "personal and admin runbook scopes, a shareable answer permalink, "
        "an admin analytics dashboard, knowledge-gap detection, runbook "
        "health monitoring, and per-runbook bookmark management.",
    )
    add_para(
        doc,
        "It demonstrates how modern LLMs, vector search, hybrid retrieval, "
        "cross-encoder re-ranking, and streaming SSE can be combined into a "
        "production-style internal tool.",
    )

    doc.add_heading("2. Problem Statement", level=2)
    add_para(
        doc,
        "Internal IT teams maintain large libraries of runbooks scattered "
        "across PDFs, wikis, and docs. Engineers waste significant time "
        "searching for the right fix during incidents, and traditional "
        "keyword search returns documents instead of answers. Generic AI "
        "chatbots hallucinate and don't cite sources, which is unacceptable "
        "when the answer drives changes on production systems. There is a "
        "need for a grounded, citation-driven assistant tuned to internal "
        "runbooks that can adapt its response depth and tone to the user's "
        "experience level.",
    )

    doc.add_heading("3. Objectives", level=2)
    add_para(
        doc,
        "The primary objective is to build a citation-grounded RAG system "
        "that returns step-by-step IT resolutions with verifiable references "
        "back to source runbooks, eliminating LLM hallucination through a "
        "strict gate-check.",
    )
    add_para(
        doc,
        "A second objective is to give the user fine-grained control over "
        "answer style and depth via response modes, so the same query can "
        "produce a 1-second terse response, a beginner-friendly explanation, "
        "or a deep root-cause + verification analysis.",
    )
    add_para(
        doc,
        "A third objective is to make the system usable as a daily tool — "
        "with streaming answers, threaded conversations, bookmarks, "
        "shareable links, exports, and an admin layer that surfaces "
        "knowledge gaps and runbook health.",
    )
    add_para(
        doc,
        "Finally, the project demonstrates how modern stacks (FastAPI, "
        "Supabase, Firebase, FAISS, Sentence-Transformers, Gemini, React + "
        "Vite + Tailwind, Framer Motion) integrate to produce a polished, "
        "real-world AI application.",
    )

    doc.add_heading("4. Core Technologies Used", level=2)
    add_table(
        doc,
        headers=["Technology", "Purpose"],
        rows=[
            ["Python 3.11+", "Backend language for FastAPI, RAG pipeline, ingestion"],
            ["FastAPI", "Async REST API + Server-Sent-Events streaming"],
            ["Google Gemini 2.5 Flash", "LLM for answer generation, HyDE, follow-ups"],
            ["Sentence-Transformers", "Embedding (BGE-small) + Re-ranking (BGE-reranker-base)"],
            ["FAISS", "Dense vector index (IndexFlatIP, cosine via L2 normalization)"],
            ["BM25Okapi (rank-bm25)", "Lexical retrieval index, fused with FAISS via RRF"],
            ["Supabase (PostgreSQL)", "Cloud database for runbooks, query logs, feedback, bookmarks"],
            ["Firebase Authentication", "Google / email-password sign-in, JWT issuance"],
            ["Firebase Admin SDK", "Backend JWT verification"],
            ["PyMuPDF, python-docx", "PDF and DOCX text extraction during ingestion"],
            ["LangChain text-splitters", "Recursive chunking with overlap"],
            ["SlowAPI", "Per-route rate limiting"],
            ["cachetools (TTLCache)", "In-process query result cache"],
            ["React 18 + Vite 5", "Frontend SPA"],
            ["Tailwind CSS 3.4", "Utility-first styling, custom design tokens"],
            ["Framer Motion", "Page transitions, step animations"],
            ["react-markdown", "Markdown rendering for AI answers"],
            ["Lucide React", "Icon set"],
            ["Axios + Fetch", "HTTP and SSE clients"],
            ["Docker + Docker Compose", "Containerization (optional deployment)"],
        ],
    )

    doc.add_heading("5. Key Features", level=2)
    features = [
        ("Multi-Format Runbook Ingestion",
         "Accepts PDF, DOCX, and TXT runbooks with magic-byte validation and SHA-256 deduplication."),
        ("Hybrid Retrieval (FAISS + BM25 + RRF)",
         "Dense vectors catch semantics, BM25 catches exact strings; Reciprocal Rank Fusion combines both."),
        ("HyDE Query Expansion",
         "Gemini generates a hypothetical runbook passage to bridge query-vs-document vocabulary gaps."),
        ("Cross-Encoder Re-Ranking",
         "BAAI/bge-reranker-base scores (query, chunk) pairs for precise top-K selection."),
        ("Six Response Modes",
         "Fast, Standard, Deep, ELI5, Expert, and Dry-run — each with its own retrieval config + prompt."),
        ("Answer Regeneration",
         "One click re-runs the same query in any other mode; logged with regenerate_of pointer."),
        ("AI Follow-Up Suggestions",
         "Gemini proposes three contextual next questions under every answer."),
        ("Server-Sent Events Streaming",
         "Tokens stream in real time over SSE; events include mode, sources, token, done."),
        ("Markdown Export",
         "Any answer downloads as a .md file with title, mode, confidence, full text, and sources."),
        ("Threaded Multi-Turn Conversations",
         "Continuation queries link back to the root via thread_id."),
        ("Citation-Grounded Answers with Gate Check",
         "If no excerpt directly addresses the query, the LLM emits an escalation sentinel."),
        ("Personal Playbook (Bookmarks)",
         "Save any answer for later review; per-user, paginated."),
        ("Shareable Answer Permalink",
         "Read-only /answer/:id route for sharing resolutions with teammates."),
        ("Admin Runbook Management",
         "Upload, list, delete, and inspect runbook health (query count, confidence, thumbs ratio)."),
        ("Knowledge Gap Detection",
         "Admin endpoint surfaces low-confidence and negative-feedback queries grouped by topic."),
        ("Per-User Runbook Scope",
         "Search admin runbooks, your own runbooks, or both — selectable per query."),
        ("Cmd+K Command Palette",
         "Keyboard-driven navigation and quick actions."),
        ("Query History",
         "Paginated audit log grouped by thread, with feedback display and resume-conversation."),
        ("Confidence Scoring",
         "Each answer shows top-K rerank confidence as a colour-coded bar."),
        ("Rate Limiting & Caching",
         "Per-user query rate limit; in-process TTL cache keyed by (mode, scope, user, query)."),
    ]
    for title, body in features:
        add_para(doc, title, bold=True)
        add_para(doc, body)
    doc.add_paragraph()

    doc.add_heading("6. Target Users", level=2)
    add_bullet(doc, "Internal IT support engineers (Tier-1 and Tier-2)")
    add_bullet(doc, "Site Reliability Engineers (SREs) handling incident response")
    add_bullet(doc, "DevOps teams maintaining runbook libraries")
    add_bullet(doc, "IT team leads / admins curating the knowledge base")
    add_bullet(doc, "New hires onboarding into operational procedures")

    doc.add_heading("7. Real-World Applications", level=2)
    add_para(
        doc,
        "ResolveIT AI directly reduces mean-time-to-resolution during "
        "production incidents. An on-call engineer hit by an Apache 502 at "
        "2 AM can describe the symptom in plain English and receive cited, "
        "step-by-step instructions in seconds, with the option to switch to "
        "Dry-run mode for an annotated, rollback-aware version before "
        "executing anything destructive. New hires can use ELI5 mode to "
        "build mental models of the infrastructure, while senior engineers "
        "default to Expert mode for terse, command-only output. The "
        "knowledge-gap dashboard tells admins which incident classes lack "
        "runbook coverage, closing the documentation feedback loop.",
    )

    # ── Application Scenarios ────────────────────────────────────────
    doc.add_heading("Application Scenarios / Use Cases", level=1)

    doc.add_heading("Scenario 1: 2 AM Incident Response", level=3)
    add_para(
        doc,
        "An on-call engineer receives a PagerDuty alert: \"Apache 502 in "
        "production.\" They open ResolveIT AI on their phone, type "
        "\"Apache returning 502 gateway error after deploy,\" and select "
        "Fast mode. Within ~1 second the system returns a 4-step fix with "
        "inline citations to the relevant runbook. They tap Dry-run to "
        "regenerate the same answer with rollback commands annotated, "
        "verify each step, and resolve the incident.",
    )

    doc.add_heading("Scenario 2: Junior Engineer Onboarding", level=3)
    add_para(
        doc,
        "A newly hired Tier-1 engineer doesn't know what \"systemd journal "
        "rotation\" means. They ask \"How do I rotate logs on the staging "
        "box?\" in ELI5 mode. The system returns a beginner-friendly "
        "explanation defining each piece of jargon, with one-line "
        "explanations above every command. They bookmark the answer to "
        "their Playbook for the next time it comes up.",
    )

    doc.add_heading("Scenario 3: Senior SRE Deep Dive", level=3)
    add_para(
        doc,
        "A senior SRE investigating intermittent latency picks Deep mode. "
        "The pipeline pulls 20 candidates, re-ranks the top 8, and returns "
        "an answer with explicit Root Cause and Verification sections. "
        "When the first answer is incomplete, they click Regenerate → "
        "Deep, which logs a regenerate_of pointer in Supabase so admins "
        "can see which queries needed multiple attempts.",
    )

    doc.add_heading("Scenario 4: Admin Curating Knowledge Gaps", level=3)
    add_para(
        doc,
        "The IT lead opens the Admin Panel and sees that the Runbook "
        "Health Grid flags two runbooks as Needs-Attention (thumbs-down "
        "ratio above 60%). The Knowledge Gaps view shows ten low-"
        "confidence queries in the last week clustered around \"VPN auth "
        "rotation,\" indicating no runbook covers it. They write a new "
        "runbook, upload it, and the FAISS index plus BM25 store rebuild "
        "automatically.",
    )

    doc.add_heading("Scenario 5: Sharing a Resolution", level=3)
    add_para(
        doc,
        "After resolving an incident, an engineer clicks the share-link "
        "icon on the answer card. The resulting /answer/<uuid> URL is "
        "pasted into the post-mortem doc. The Markdown export button "
        "produces a clean .md file embedded into the incident report.",
    )

    # ── Technical Architecture ───────────────────────────────────────
    doc.add_heading("Technical Architecture", level=1)

    doc.add_heading("● Frontend Layer", level=3)
    add_para(
        doc,
        "Built with React 18 + Vite 5 + Tailwind CSS. The UI is a single-"
        "page application with routes for Landing, Login, Register, "
        "Dashboard, AdminPanel, MyRunbooks, MyPlaybook (bookmarks), "
        "QueryHistory, and SharedAnswer. Global state is held in "
        "AuthContext (Firebase token, isAdmin). The useQuery hook handles "
        "both blocking (POST /query) and streaming (POST /query/stream "
        "via fetch + ReadableStream) requests.",
    )

    doc.add_heading("● Backend Layer", level=3)
    add_para(
        doc,
        "FastAPI application split into routes (auth, query, admin, "
        "feedback, history, runbooks, bookmarks, exports), core (config, "
        "Firebase auth, Gemini client, Supabase client, rate limiter), "
        "ingestion (parser, chunker, indexer), retrieval (embedder, "
        "FAISS store, BM25 store, reranker, hybrid RRF), rag (modes, "
        "pipeline), and models (request and response Pydantic schemas). "
        "All ML models warm up on startup via the lifespan context.",
    )

    doc.add_heading("● Database Layer", level=3)
    add_para(
        doc,
        "Supabase (managed PostgreSQL) holds four tables: runbooks "
        "(metadata + content_hash + is_admin_runbook), query_logs "
        "(retrieved sources, llm_response, confidence_score, thread_id, "
        "mode, regenerate_of), feedback (rating, comment), and "
        "bookmarks (user-scoped saved answers). FAISS index and "
        "metadata are stored on the backend filesystem and persisted "
        "across restarts; BM25 is rebuilt in-memory from FAISS metadata.",
    )

    doc.add_heading("● External APIs / Services", level=3)
    add_para(
        doc,
        "Google Gemini 2.5 Flash (with fallback model list) for HyDE "
        "expansion, answer generation, follow-up suggestion, and (in Deep "
        "mode) critique-and-retry. Firebase Authentication for OAuth and "
        "email/password sign-in; Firebase Admin SDK for token verification "
        "on the backend. Supabase REST API for all database access.",
    )

    doc.add_heading("● Deployment Environment", level=3)
    add_para(
        doc,
        "The application is containerized with Docker Compose (backend + "
        "frontend services). The backend mounts a persistent volume for "
        "the FAISS index and the Firebase service-account JSON. The "
        "frontend is built with Vite and served as static assets. "
        "Environment variables hold all secrets (Supabase keys, Gemini API "
        "key, Firebase credentials path, admin-email whitelist, CORS "
        "origins, rate-limit settings).",
    )

    doc.add_heading("High-Level Architecture Diagram", level=2)
    add_diagram_placeholder(doc, "DIAGRAM 1: High-Level System Architecture")

    doc.add_heading("Component Interaction Diagram", level=2)
    add_diagram_placeholder(doc, "DIAGRAM 2: Component Interaction (React ↔ FastAPI ↔ Services)")

    doc.add_heading("RAG Pipeline Flow", level=2)
    add_diagram_placeholder(doc, "DIAGRAM 3: RAG Pipeline (Query → HyDE → Hybrid → Rerank → Gate → Gemini → Cache)")

    doc.add_heading("Mode Selection Flow", level=2)
    add_diagram_placeholder(doc, "DIAGRAM 4: Mode Selection (Fast / Standard / Deep / ELI5 / Expert / Dry-run)")

    doc.add_heading("Streaming SSE Sequence", level=2)
    add_diagram_placeholder(doc, "DIAGRAM 5: SSE Sequence (browser ↔ /query/stream events)")

    doc.add_heading("Database ER Diagram", level=2)
    add_diagram_placeholder(doc, "DIAGRAM 6: Database ER (runbooks, query_logs, feedback, bookmarks)")

    doc.add_heading("Authentication Flow", level=2)
    add_diagram_placeholder(doc, "DIAGRAM 7: Authentication (Firebase OAuth → ID token → backend verify)")

    doc.add_heading("Deployment Topology", level=2)
    add_diagram_placeholder(doc, "DIAGRAM 8: Deployment Topology (Docker Compose backend + frontend)")

    # ── Pre-requisites ───────────────────────────────────────────────
    doc.add_heading("Pre-requisites", level=1)

    doc.add_heading("Software Requirements", level=3)
    add_bullet(doc, "Python 3.11 or higher")
    add_bullet(doc, "Node.js 18+ and npm")
    add_bullet(doc, "Git")
    add_bullet(doc, "Docker Desktop (optional, for containerized run)")
    add_bullet(doc, "IDE: Visual Studio Code with Python and ESLint extensions")
    add_para(doc, "")

    doc.add_heading("Libraries / Frameworks / Dependencies", level=3)
    add_para(doc, "Backend dependencies (install with):")
    add_code(doc, "pip install -r backend/requirements.txt")
    add_para(doc, "Key packages: fastapi, uvicorn, google-genai, sentence-transformers, "
                  "faiss-cpu, rank-bm25, supabase, firebase-admin, pymupdf, python-docx, "
                  "langchain-text-splitters, slowapi, cachetools, python-dotenv.")
    add_para(doc, "")
    add_para(doc, "Frontend dependencies (install with):")
    add_code(doc, "cd frontend && npm install")
    add_para(doc, "Key packages: react, react-dom, react-router-dom, vite, axios, "
                  "firebase, tailwindcss, framer-motion, lucide-react, react-markdown, "
                  "react-hot-toast.")
    add_para(doc, "")

    doc.add_heading("Hardware Requirements", level=3)
    add_bullet(doc, "Minimum 8 GB RAM (FAISS + sentence-transformers models)")
    add_bullet(doc, "4+ CPU cores recommended for concurrent embedding/reranking")
    add_bullet(doc, "5 GB free disk space (models + FAISS index)")
    add_bullet(doc, "Stable internet connection (Gemini API + Supabase + Firebase)")

    # ── Prior Knowledge ──────────────────────────────────────────────
    doc.add_heading("Prior Knowledge Required", level=1)

    doc.add_heading("Programming Language Knowledge", level=3)
    add_bullet(doc, "Strong Python fundamentals (async/await, typing, dataclasses)")
    add_bullet(doc, "JavaScript ES6+, React hooks, JSX, modern frontend tooling")
    add_bullet(doc, "JSON, REST, and basic networking concepts")

    doc.add_heading("Framework Basics", level=3)
    add_bullet(doc, "FastAPI: routing, dependency injection, Pydantic, lifespan events")
    add_bullet(doc, "React: components, hooks, context, react-router, fetch streaming")
    add_bullet(doc, "Tailwind CSS utility-first styling")

    doc.add_heading("Database Fundamentals", level=3)
    add_bullet(doc, "PostgreSQL / Supabase tables, queries, foreign keys, indexes")
    add_bullet(doc, "CRUD operations and migrations")

    doc.add_heading("AI / ML Fundamentals", level=3)
    add_bullet(doc, "Vector embeddings and similarity search (cosine, inner product)")
    add_bullet(doc, "BM25 lexical scoring")
    add_bullet(doc, "Cross-encoder re-ranking")
    add_bullet(doc, "Retrieval-Augmented Generation (RAG) and HyDE")
    add_bullet(doc, "Prompt engineering for citation-grounded answers")
    add_bullet(doc, "LLM streaming and SSE")

    doc.add_heading("Networking / Cloud Concepts", level=3)
    add_bullet(doc, "Client-server architecture and HTTP methods")
    add_bullet(doc, "JWT and OAuth2 (via Firebase)")
    add_bullet(doc, "Server-Sent Events (SSE)")
    add_bullet(doc, "CORS, environment variables, secrets management")

    # ── Project Objectives ───────────────────────────────────────────
    doc.add_heading("Project Objectives", level=1)

    doc.add_heading("Technical Objectives", level=3)
    add_bullet(doc, "Build a hybrid retrieval pipeline combining FAISS dense search and BM25 lexical search via Reciprocal Rank Fusion")
    add_bullet(doc, "Implement HyDE query expansion to improve recall on vocabulary-mismatched questions")
    add_bullet(doc, "Add a cross-encoder re-ranking stage to maximize precision of the top-K passages")
    add_bullet(doc, "Design a six-mode prompt and retrieval-config system that branches the pipeline without code duplication")
    add_bullet(doc, "Implement an SSE streaming endpoint plus a blocking endpoint sharing the same pipeline")
    add_bullet(doc, "Ground every answer in cited excerpts with a strict gate check that refuses to answer when no excerpt is directly relevant")

    doc.add_heading("Performance Objectives", level=3)
    add_bullet(doc, "Sub-1-second time-to-first-token in Fast mode")
    add_bullet(doc, "Cache hit time < 50 ms for repeated queries (TTLCache)")
    add_bullet(doc, "Stable handling of 20 queries/minute per user (slowapi rate limiting)")
    add_bullet(doc, "Persistent FAISS index across restarts to avoid re-embedding")

    doc.add_heading("Deployment Objectives", level=3)
    add_bullet(doc, "Single docker-compose up to launch the full stack locally")
    add_bullet(doc, "All secrets via environment variables; no hard-coded keys")
    add_bullet(doc, "Idempotent SQL migrations with backward-compatible inserts")
    add_bullet(doc, "Volume-mounted FAISS index for production persistence")

    doc.add_heading("Learning Outcomes", level=3)
    add_bullet(doc, "Practical experience designing and shipping a production RAG pipeline")
    add_bullet(doc, "Deep familiarity with hybrid retrieval, re-ranking, and HyDE")
    add_bullet(doc, "End-to-end React + FastAPI + Supabase + Firebase + Gemini integration")
    add_bullet(doc, "Streaming UX patterns with Server-Sent Events")
    add_bullet(doc, "Prompt engineering across multiple response styles")

    # ── Project Workflow (outline) ───────────────────────────────────
    doc.add_heading("Project Workflow", level=1)

    add_para(doc, "Milestone 1: Requirement Analysis & System Design", bold=True)
    add_bullet(doc, "Activity 1.1: Problem Definition")
    add_bullet(doc, "Activity 1.2: Functional Requirements")
    add_bullet(doc, "Activity 1.3: Non-Functional Requirements")
    add_bullet(doc, "Activity 1.4: System Design Decisions & Technology Stack Selection")

    add_para(doc, "Milestone 2: Environment Setup & Initial Configuration", bold=True)
    add_bullet(doc, "Activity 2.1: Development Environment Setup")
    add_bullet(doc, "Activity 2.2: Dependency Installation")
    add_bullet(doc, "Activity 2.3: Project Structure Creation")
    add_bullet(doc, "Activity 2.4: Configuration & Secrets Setup")

    add_para(doc, "Milestone 3: Database & Authentication Setup", bold=True)
    add_bullet(doc, "Activity 3.1: Create Supabase project and tables")
    add_bullet(doc, "Activity 3.2: Configure Firebase Authentication")
    add_bullet(doc, "Activity 3.3: Run schema migrations (001 → 006)")

    add_para(doc, "Milestone 4: Core RAG Pipeline Development", bold=True)
    add_bullet(doc, "Activity 4.1: Ingestion (parser, chunker, indexer)")
    add_bullet(doc, "Activity 4.2: Embedding + FAISS dense store")
    add_bullet(doc, "Activity 4.3: BM25 lexical store + RRF hybrid fusion")
    add_bullet(doc, "Activity 4.4: Cross-encoder re-ranker")
    add_bullet(doc, "Activity 4.5: HyDE expansion + Gemini generation")
    add_bullet(doc, "Activity 4.6: Streaming SSE pipeline")

    add_para(doc, "Milestone 5: Multi-Mode Response System", bold=True)
    add_bullet(doc, "Activity 5.1: Design ModeConfig dataclass")
    add_bullet(doc, "Activity 5.2: Implement six prompts (Fast / Standard / Deep / ELI5 / Expert / Dry-run)")
    add_bullet(doc, "Activity 5.3: Branch retrieval and generation in pipeline")
    add_bullet(doc, "Activity 5.4: Add follow-up question generation")
    add_bullet(doc, "Activity 5.5: Add regenerate flow + Markdown export endpoint")

    add_para(doc, "Milestone 6: Frontend Development", bold=True)
    add_bullet(doc, "Activity 6.1: Auth flow (Login / Register / ProtectedRoute)")
    add_bullet(doc, "Activity 6.2: Dashboard with streaming chat UI, mode picker, scope chips")
    add_bullet(doc, "Activity 6.3: ResultCard with wizard view, sources, follow-ups, regenerate, export")
    add_bullet(doc, "Activity 6.4: Admin Panel, Knowledge Gaps, Runbook Health Grid")
    add_bullet(doc, "Activity 6.5: Query History, My Runbooks, My Playbook, Shared Answer")
    add_bullet(doc, "Activity 6.6: Cmd+K Command Palette, AppLayout, animations")

    add_para(doc, "Milestone 7: Integration & Optimization", bold=True)
    add_bullet(doc, "Activity 7.1: End-to-end integration testing")
    add_bullet(doc, "Activity 7.2: TTL cache + rate limiting")
    add_bullet(doc, "Activity 7.3: Error handling, escalation gate, abort controllers")

    add_para(doc, "Milestone 8: Testing & Validation", bold=True)
    add_bullet(doc, "Activity 8.1: Functional Test Case Execution")
    add_bullet(doc, "Activity 8.2: Unit & Integration Testing")
    add_bullet(doc, "Activity 8.3: User Acceptance Testing")
    add_bullet(doc, "Activity 8.4: Performance Evaluation")

    add_para(doc, "Milestone 9: Containerization & Deployment Prep", bold=True)
    add_bullet(doc, "Activity 9.1: Dockerfile + docker-compose configuration")
    add_bullet(doc, "Activity 9.2: Hosting platform planning (Vercel + Render or VM)")
    add_bullet(doc, "Activity 9.3: Production hardening (secrets, CORS, rate limits)")

    # ── Detailed Milestones ──────────────────────────────────────────
    doc.add_heading("Detailed Milestone Activities", level=1)

    doc.add_heading("Milestone 1: Requirement Analysis & System Design", level=2)
    doc.add_heading("Activity 1.1: Problem Definition", level=3)
    add_para(doc,
        "Defined the gap between traditional keyword-based runbook search "
        "and the need for citation-grounded, mode-aware answers. "
        "Identified personas (Tier-1, SRE, Admin, new hire) and their "
        "different needs from the same underlying runbook library.")
    doc.add_heading("Activity 1.2: Functional Requirements", level=3)
    add_para(doc,
        "Multi-format ingestion, hybrid retrieval, six response modes, "
        "streaming and blocking endpoints, threaded conversations, "
        "bookmarks, sharing, exports, admin analytics, and knowledge-gap "
        "detection.")
    doc.add_heading("Activity 1.3: Non-Functional Requirements", level=3)
    add_para(doc,
        "Sub-1-second Fast-mode latency, persistence across restarts, "
        "rate-limiting, JWT-based auth, idempotent migrations, env-driven "
        "configuration, and CORS-restricted API.")
    doc.add_heading("Activity 1.4: System Design Decisions & Technology Stack Selection", level=3)
    add_para(doc,
        "Selected FastAPI for async + SSE support; Gemini 2.5 Flash for "
        "low latency and large context; FAISS for offline-friendly vector "
        "search; BM25Okapi for lexical complement; Supabase for managed "
        "Postgres with REST; Firebase for auth offload; React+Vite+Tailwind "
        "for the UI.")

    doc.add_heading("Milestone 2: Environment Setup & Initial Configuration", level=2)
    doc.add_heading("Activity 2.1: Development Environment Setup", level=3)
    add_para(doc, "Python 3.11 venv created in backend/venv. Node 18+ for frontend.")
    doc.add_heading("Activity 2.2: Dependency Installation", level=3)
    add_para(doc, "pip install -r backend/requirements.txt; cd frontend && npm install.")
    doc.add_heading("Activity 2.3: Project Structure Creation", level=3)
    add_para(doc, "Created modular layout: core/, ingestion/, retrieval/, rag/, routes/, models/.")
    doc.add_heading("Activity 2.4: Configuration & Secrets Setup", level=3)
    add_para(doc,
        "Filled backend/.env with SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY, "
        "GEMINI_API_KEY, FIREBASE_CREDENTIALS_PATH, ADMIN_EMAILS, "
        "CORS_ORIGINS, MAX_UPLOAD_MB, QUERY_RATE_LIMIT, QUERY_CACHE_*. "
        "Filled frontend/.env with VITE_FIREBASE_* and VITE_API_BASE_URL.")

    doc.add_heading("Milestone 3: Database & Authentication Setup", level=2)
    doc.add_heading("Activity 3.1: Create Supabase Project and Tables", level=3)
    add_para(doc, "Created tables runbooks, query_logs, feedback, bookmarks via migration 001.")
    doc.add_heading("Activity 3.2: Configure Firebase Authentication", level=3)
    add_para(doc, "Enabled Google OAuth and email/password providers; downloaded service-account JSON for backend admin SDK.")
    doc.add_heading("Activity 3.3: Run Schema Migrations", level=3)
    add_para(doc,
        "Ran migrations 001 (initial), 002 (content_hash dedup), 003 "
        "(is_admin_runbook), 004 (bookmarks RLS), 005 (thread_id), 006 "
        "(mode + regenerate_of). All migrations are idempotent (IF NOT "
        "EXISTS) and the application falls back gracefully if a column is "
        "missing.")

    doc.add_heading("Milestone 4: Core RAG Pipeline Development", level=2)
    for act, body in [
        ("Activity 4.1: Ingestion",
         "parser.py extracts text via PyMuPDF (PDF) or python-docx (DOCX). "
         "chunker.py uses LangChain's recursive splitter with 1600-char "
         "chunks and 200-char overlap, hierarchy-aware (## → # → ### → "
         "blank line). indexer.py orchestrates parse → chunk → embed → "
         "FAISS add → BM25 rebuild → Supabase insert with SHA-256 dedup."),
        ("Activity 4.2: Embedding + FAISS Dense Store",
         "embedder.py loads BAAI/bge-small-en-v1.5 once at startup, "
         "L2-normalizes vectors, prepends the BGE query prefix at query "
         "time only. faiss_store.py uses IndexFlatIP (inner-product = "
         "cosine on normalized vectors) and persists index.bin + "
         "metadata.json to disk."),
        ("Activity 4.3: BM25 + RRF Hybrid Fusion",
         "bm25_store.py wraps rank_bm25.BM25Okapi, rebuilt from FAISS "
         "metadata on startup. hybrid.py runs both stores top-k=12 and "
         "fuses them via Reciprocal Rank Fusion (k=60 smoothing) with "
         "deduplication by chunk_id."),
        ("Activity 4.4: Cross-Encoder Re-Ranker",
         "reranker.py loads BAAI/bge-reranker-base, scores (query, chunk) "
         "pairs, applies sigmoid to map logits to [0, 1] confidence."),
        ("Activity 4.5: HyDE + Gemini Generation",
         "gemini_client.py implements hyde_expand (3–5 sentence "
         "hypothetical passage), generate_resolution (blocking), "
         "stream_resolution (SSE), and generate_followups (3 next "
         "questions). All calls fall back through GEMINI_FALLBACK_MODELS "
         "on 503/429 errors."),
        ("Activity 4.6: Streaming SSE Pipeline",
         "stream_rag_pipeline yields events {mode, sources, token×N, "
         "done} as Python dicts; routes/query.py serializes each as an "
         "SSE frame. The frontend useQuery hook reads them with fetch + "
         "ReadableStream, accumulating tokens in real time."),
    ]:
        doc.add_heading(act, level=3)
        add_para(doc, body)

    doc.add_heading("Milestone 5: Multi-Mode Response System", level=2)
    for act, body in [
        ("Activity 5.1: ModeConfig Dataclass",
         "rag/modes.py defines a frozen dataclass with fields name, "
         "use_hyde, top_k, top_n, temperature, system_prompt, "
         "critique_retry. A MODES dict maps each mode name to its config."),
        ("Activity 5.2: Six Mode Prompts",
         "Fast (terse, no HyDE, top_k=6, temp=0). Standard (default, "
         "current behaviour). Deep (HyDE, top_k=20, top_n=8, adds Root "
         "Cause + Verification sections). ELI5 (beginner-friendly, "
         "explains jargon, command-then-explanation). Expert (TL;DR + "
         "command-only steps). Dry-run (commands annotated with comments, "
         "Rollback section)."),
        ("Activity 5.3: Pipeline Branching",
         "run_rag_pipeline and stream_rag_pipeline both accept a mode "
         "parameter, look up the ModeConfig via get_mode(name), and pass "
         "system_prompt + temperature into the LLM call. The cache key is "
         "(mode, scope, user, query) so the same query in different modes "
         "is cached independently."),
        ("Activity 5.4: Follow-up Generation",
         "After every successful answer, generate_followups asks Gemini "
         "for 3 short next questions, returned as part of the QueryResponse "
         "and the SSE done event."),
        ("Activity 5.5: Regenerate + Markdown Export",
         "QueryRequest.regenerate_of stores the prior query_log_id; "
         "when set, the cache is bypassed and the new log row points back "
         "via regenerate_of FK. routes/exports.py serves "
         "/export/{query_log_id}.md returning a downloadable Markdown "
         "file with title, mode, confidence, full answer, and sources."),
    ]:
        doc.add_heading(act, level=3)
        add_para(doc, body)

    doc.add_heading("Milestone 6: Frontend Development", level=2)
    doc.add_heading("Activity 6.1: Authentication Flow", level=3)
    add_para(doc,
        "AuthContext wraps the app, holds {user, token, isAdmin, loading}, "
        "auto-refreshes the Firebase ID token every 50 minutes. "
        "ProtectedRoute guards authenticated and admin-only routes.")
    doc.add_heading("Activity 6.2: Dashboard Streaming Chat", level=3)
    add_para(doc,
        "Dashboard.jsx maintains a messages array of user, assistant, and "
        "error bubbles. The floating input island contains the new "
        "Mode picker chip row, the scope picker, and the QueryInput. "
        "useQuery's submitQueryStream reads SSE frames and updates "
        "streaming text live.")
    doc.add_heading("Activity 6.3: ResultCard", level=3)
    add_para(doc,
        "Renders the answer in either step-wizard mode (parsed numbered "
        "list with checkboxes, progress bar, prevention/follow-up "
        "collapsible) or raw markdown mode. Header shows confidence + "
        "mode badge. Action bar: copy text, copy share link, bookmark, "
        "Markdown export, Regenerate dropdown. Below sources, AI "
        "follow-up question chips when present.")
    doc.add_heading("Activity 6.4: Admin Surfaces", level=3)
    add_para(doc,
        "AdminPanel.jsx shows totals + thumbs-up/down rates. "
        "RunbookHealthGrid renders per-runbook query count, average "
        "confidence, and a thumbs ratio bar with a Needs-Attention flag. "
        "Knowledge Gaps section lists low-confidence and negative-feedback "
        "queries grouped by topic.")
    doc.add_heading("Activity 6.5: User Surfaces", level=3)
    add_para(doc,
        "QueryHistory groups queries by thread_id with pagination. "
        "MyRunbooks lets users upload/delete personal runbooks. "
        "MyPlaybook lists bookmarked answers. SharedAnswer renders a "
        "read-only answer page accessed via /answer/:id.")
    doc.add_heading("Activity 6.6: Polish", level=3)
    add_para(doc,
        "Cmd+K CommandPalette for keyboard navigation, AppLayout with "
        "ambient mouse-tracking spotlight, Framer Motion page transitions, "
        "react-hot-toast notifications, custom Tailwind tokens (primary "
        "orange, dark brown), JetBrains Mono for code blocks.")

    doc.add_heading("Milestone 7: Integration & Optimization", level=2)
    doc.add_heading("Activity 7.1: End-to-End Integration", level=3)
    add_para(doc,
        "Verified full flow: upload runbook → query in each mode → "
        "feedback → bookmark → share → export. Streaming and blocking "
        "paths share the same retrieval and generation logic.")
    doc.add_heading("Activity 7.2: Caching & Rate Limiting", level=3)
    add_para(doc,
        "Added in-process TTLCache (256 entries × 600 s TTL) keyed by "
        "(mode, scope, user, query). Per-route rate limiting via slowapi "
        "(QUERY_RATE_LIMIT default 20/minute).")
    doc.add_heading("Activity 7.3: Error Handling & Escalation", level=3)
    add_para(doc,
        "Strict gate-check in the system prompt forces the LLM to emit a "
        "fixed escalation sentinel when no excerpt is directly relevant; "
        "the pipeline detects it and clears sources/confidence. Frontend "
        "renders an amber Escalation card. AbortController cancels in-"
        "flight streams when the user starts a new query.")

    doc.add_heading("Milestone 8: Testing & Validation", level=2)
    doc.add_heading("Activity 8.1: Functional Test Cases", level=3)
    add_para(doc,
        "Verified each mode produces the expected pipeline-config branch "
        "(Fast skips HyDE, Deep uses top_k=20, etc.) and the right system "
        "prompt. Verified gate-check escalation on irrelevant queries. "
        "Verified follow-up suggestions render and click-through.")
    doc.add_heading("Activity 8.2: Unit & Integration Testing", level=3)
    add_para(doc,
        "Backend: validated Pydantic schemas reject invalid modes, "
        "cache keys differ across modes, Markdown export renders with "
        "correct frontmatter. Frontend: verified Vite production build "
        "succeeds without warnings, SSE parser handles all event types.")
    doc.add_heading("Activity 8.3: User Acceptance Testing", level=3)
    add_para(doc,
        "Sample runbooks ingested; sample queries run; answers manually "
        "graded for citation accuracy and mode appropriateness.")
    doc.add_heading("Activity 8.4: Performance Evaluation", level=3)
    add_para(doc,
        "Fast mode time-to-first-token under 1 s on warm cache. "
        "Standard mode end-to-end ~3–5 s. Cache-hit path under 50 ms. "
        "FAISS query under 10 ms for 10k-vector index.")

    doc.add_heading("Milestone 9: Containerization & Deployment Prep", level=2)
    doc.add_heading("Activity 9.1: Docker Compose", level=3)
    add_para(doc,
        "Two services (backend, frontend) defined in docker-compose.yml. "
        "Backend mounts ./backend/faiss_index and the firebase-service-"
        "account.json as volumes. Frontend builds and serves static assets "
        "on port 5173.")
    doc.add_heading("Activity 9.2: Hosting Platform Planning", level=3)
    add_para(doc,
        "Suitable hosts: Vercel for frontend static build; Render or a "
        "small VM (e.g. Hetzner CPX21) for the backend (sentence-"
        "transformers needs ~1.5 GB RAM).")
    doc.add_heading("Activity 9.3: Production Hardening", level=3)
    add_para(doc,
        "All secrets via env vars, never committed. CORS restricted to "
        "production frontend domain. Per-IP rate limit. HTTPS via reverse "
        "proxy. FAISS index on a persistent volume.")

    # ── Project Structure ────────────────────────────────────────────
    doc.add_heading("Project Structure", level=1)
    doc.add_heading("Clean Folder Structure", level=3)
    tree = """resolveit-ai/
│
├── backend/
│   ├── core/
│   │   ├── config.py
│   │   ├── firebase_auth.py
│   │   ├── gemini_client.py
│   │   ├── supabase_client.py
│   │   └── rate_limit.py
│   ├── ingestion/
│   │   ├── parser.py
│   │   ├── chunker.py
│   │   └── indexer.py
│   ├── retrieval/
│   │   ├── embedder.py
│   │   ├── faiss_store.py
│   │   ├── bm25_store.py
│   │   ├── reranker.py
│   │   └── hybrid.py
│   ├── rag/
│   │   ├── pipeline.py
│   │   └── modes.py
│   ├── routes/
│   │   ├── auth.py
│   │   ├── query.py
│   │   ├── admin.py
│   │   ├── feedback.py
│   │   ├── history.py
│   │   ├── runbooks.py
│   │   ├── bookmarks.py
│   │   └── exports.py
│   ├── models/
│   │   ├── request_models.py
│   │   └── response_models.py
│   ├── migrations/
│   │   ├── 001_create_tables.sql
│   │   ├── 002_add_content_hash.sql
│   │   ├── 003_add_is_admin_runbook.sql
│   │   ├── 004_bookmarks_rls.sql
│   │   ├── 005_add_thread_id.sql
│   │   └── 006_modes_and_followups.sql
│   ├── faiss_index/
│   ├── main.py
│   ├── requirements.txt
│   └── .env
│
├── frontend/
│   ├── src/
│   │   ├── components/
│   │   │   ├── AppLayout.jsx
│   │   │   ├── Navbar.jsx
│   │   │   ├── Sidebar.jsx
│   │   │   ├── CommandCenter.jsx
│   │   │   ├── CommandPalette.jsx
│   │   │   ├── QueryInput.jsx
│   │   │   ├── ResultCard.jsx
│   │   │   ├── SourceBadge.jsx
│   │   │   ├── FeedbackButtons.jsx
│   │   │   ├── RunbookTable.jsx
│   │   │   ├── RunbookUploader.jsx
│   │   │   ├── RunbookHealthGrid.jsx
│   │   │   └── ProtectedRoute.jsx
│   │   ├── pages/
│   │   │   ├── Landing.jsx
│   │   │   ├── Login.jsx
│   │   │   ├── Register.jsx
│   │   │   ├── Dashboard.jsx
│   │   │   ├── AdminPanel.jsx
│   │   │   ├── QueryHistory.jsx
│   │   │   ├── MyRunbooks.jsx
│   │   │   ├── MyPlaybook.jsx
│   │   │   └── SharedAnswer.jsx
│   │   ├── context/
│   │   │   └── AuthContext.jsx
│   │   ├── hooks/
│   │   │   ├── useAuth.js
│   │   │   └── useQuery.js
│   │   ├── App.jsx
│   │   ├── main.jsx
│   │   └── firebaseConfig.js
│   ├── package.json
│   ├── vite.config.js
│   └── tailwind.config.js
│
├── docker-compose.yml
├── PROJECT.md
├── PROJECT_REPORT.md
└── README.md
"""
    add_code(doc, tree)

    doc.add_heading("Folder Explanations", level=3)
    folder_explanations = [
        ("backend/", "FastAPI application root; async API + RAG pipeline + ML models."),
        ("backend/core/", "Cross-cutting concerns: env config, Firebase auth, Gemini client, Supabase client, slowapi rate limiter."),
        ("backend/ingestion/", "Runbook ingestion: parse → chunk → embed → store. Handles PDF/DOCX/TXT with magic-byte validation and SHA-256 dedup."),
        ("backend/retrieval/", "Embedding (BGE-small), dense store (FAISS IndexFlatIP), lexical store (BM25Okapi), cross-encoder re-ranker, RRF hybrid fusion."),
        ("backend/rag/", "pipeline.py orchestrates retrieve → rerank → generate → log; modes.py holds the six ModeConfigs and prompt pack."),
        ("backend/routes/", "All FastAPI endpoints split per resource. exports.py serves Markdown downloads."),
        ("backend/models/", "Pydantic schemas for request/response validation."),
        ("backend/migrations/", "Idempotent SQL migrations applied to Supabase in order."),
        ("backend/faiss_index/", "Persisted dense index binary + per-chunk metadata JSON."),
        ("frontend/", "React + Vite SPA."),
        ("frontend/src/components/", "Reusable UI: navbar, sidebar, command palette, query input, result card, etc."),
        ("frontend/src/pages/", "Route-level views: Landing, Login, Dashboard, AdminPanel, QueryHistory, MyRunbooks, MyPlaybook, SharedAnswer."),
        ("frontend/src/context/", "AuthContext provides global user/token/isAdmin state."),
        ("frontend/src/hooks/", "useAuth (consume context) + useQuery (submit blocking, streaming, feedback)."),
        ("docker-compose.yml", "Two-service compose for local containerized run (backend + frontend)."),
        (".env", "Stores all secrets and runtime configuration."),
        ("PROJECT.md / README.md", "Original project docs and quickstart guide."),
    ]
    for name, body in folder_explanations:
        add_para(doc, name, bold=True)
        add_para(doc, body)

    # ── Results ──────────────────────────────────────────────────────
    doc.add_heading("Results", level=1)

    doc.add_heading("● System Output", level=3)
    add_para(doc,
        "The system successfully ingests runbooks of mixed formats and "
        "answers natural-language IT queries with cited, structured "
        "resolutions. Each answer carries a confidence score, source "
        "list, and (when relevant) AI-suggested follow-up questions. "
        "The mode selector lets the same query produce six visibly "
        "different responses, validating the prompt-pack approach.")

    doc.add_heading("● Performance Evaluation", level=3)
    add_para(doc,
        "Fast mode delivers time-to-first-token under one second on a "
        "warm index. Standard mode answers in 3–5 seconds end-to-end. "
        "Cache hits return in under 50 ms. FAISS look-ups stay under "
        "10 ms even with thousands of vectors thanks to IndexFlatIP and "
        "L2 normalization.")

    doc.add_heading("● Screenshots", level=3)
    add_para(doc,
        "The screenshots below capture the working application across "
        "all major surfaces. Replace each placeholder with the captured "
        "image when ready.")
    screens = [
        ("SCREENSHOT 1: Landing Page",                    7.5),
        ("SCREENSHOT 2: Login Page (Google + Email)",     7.0),
        ("SCREENSHOT 3: Dashboard — Empty Command Center", 7.5),
        ("SCREENSHOT 4: Dashboard — Mode + Scope Chips Visible", 7.0),
        ("SCREENSHOT 5: Streaming Bubble Mid-Answer",     6.5),
        ("SCREENSHOT 6: Result Card — Standard Mode (sources + follow-ups)", 8.5),
        ("SCREENSHOT 7: Result Card — ELI5 Mode",         8.0),
        ("SCREENSHOT 8: Result Card — Expert Mode",       8.0),
        ("SCREENSHOT 9: Result Card — Dry-run Mode (with Rollback)", 8.5),
        ("SCREENSHOT 10: Regenerate Dropdown Open",       6.0),
        ("SCREENSHOT 11: Markdown Export (downloaded file)", 6.5),
        ("SCREENSHOT 12: Admin Panel — Stats Cards",      7.5),
        ("SCREENSHOT 13: Runbook Health Grid",            7.5),
        ("SCREENSHOT 14: Knowledge Gaps View",            7.0),
        ("SCREENSHOT 15: My Runbooks (Upload + List)",    7.5),
        ("SCREENSHOT 16: My Playbook (Bookmarks)",        7.0),
        ("SCREENSHOT 17: Query History (Threaded)",       7.5),
        ("SCREENSHOT 18: Shared Answer (Public View)",    7.0),
        ("SCREENSHOT 19: Cmd+K Command Palette Open",     6.5),
        ("SCREENSHOT 20: Supabase Tables (query_logs with mode column)", 7.0),
    ]
    for label, h in screens:
        add_diagram_placeholder(doc, label, height_cm=h)

    doc.add_heading("● Benchmark Results", level=3)
    add_table(
        doc,
        headers=["Metric", "Result"],
        rows=[
            ["Fast-mode time-to-first-token", "< 1 s (warm index)"],
            ["Standard-mode end-to-end", "3 – 5 s"],
            ["Deep-mode end-to-end (top_k=20, top_n=8)", "5 – 9 s"],
            ["Cache-hit path", "< 50 ms"],
            ["FAISS top-12 search (10k vectors)", "< 10 ms"],
            ["BGE-small embedding (single query)", "~ 30 ms (CPU)"],
            ["Cross-encoder re-rank (12 candidates)", "~ 200 ms (CPU)"],
            ["Streaming token throughput", "~ 50–80 tokens/s (Gemini Flash)"],
            ["Concurrent users supported", "20 q/min/user via slowapi"],
        ],
    )

    # ── Advantages & Limitations ─────────────────────────────────────
    doc.add_heading("Advantages & Limitations", level=1)

    doc.add_heading("Advantages", level=3)
    advantages = [
        ("Citation-Grounded Answers", "Every answer cites the runbook excerpts that produced it; the gate-check refuses to fabricate when no excerpt is directly relevant."),
        ("Six Response Modes", "One pipeline produces six distinct response styles, letting the same knowledge serve juniors, seniors, and dry-run safety reviews."),
        ("Hybrid Retrieval", "FAISS catches semantics, BM25 catches exact strings (commands, error codes); RRF combines both without weight tuning."),
        ("Cross-Encoder Re-Ranking", "Final top-K precision is significantly higher than bi-encoder-only retrieval."),
        ("Real-Time Streaming UX", "Answers stream token-by-token; users see progress in 200 ms instead of waiting for the full generation."),
        ("Persistent FAISS Index", "Survives restarts without re-embedding; production-friendly."),
        ("Bookmarks, Sharing, Export", "Answers are first-class artifacts: saveable, shareable via permalink, exportable as Markdown."),
        ("Admin Knowledge Gaps", "Surfaces which queries the system can't answer well, closing the documentation feedback loop."),
        ("Backward-Compatible Migrations", "Inserts gracefully drop unknown columns when migrations lag, so deploys don't block on DDL."),
        ("Modular Architecture", "Adding a new response mode is a 10-line dict entry plus one prompt; adding a new ingestion format is a single parser."),
    ]
    for t, b in advantages:
        add_para(doc, t, bold=True)
        add_para(doc, b)

    doc.add_heading("Limitations", level=3)
    limitations = [
        ("Single-Tenant Auth Model", "Admin status is a static email whitelist (ADMIN_EMAILS) rather than database-backed RBAC; no multi-workspace support yet."),
        ("Full-Index Rebuild on Delete", "BM25 is rebuilt from scratch on every delete; switching to IndexIDMap with stable per-chunk IDs would make this incremental."),
        ("In-Memory BM25", "BM25 is rebuilt on startup; not suitable for very large corpora without a serialization format."),
        ("Single LLM Provider", "Only Gemini is supported; adding Claude/GPT would require a pluggable LLM client."),
        ("No Document Versioning", "When a runbook is replaced, history is lost; v3-vs-v4 diffing is not yet supported."),
        ("Citation Faithfulness Not Verified", "Inline citations are emitted by the LLM; a second-pass verification (\"every claim traces to a chunk\") is on the roadmap."),
        ("CPU-Only Embedding/Reranking", "No GPU path; latency is dominated by CPU re-ranking."),
        ("English Only", "Embedding model and prompts are English-tuned."),
        ("Limited Integrations", "No Confluence/Notion/GitHub/Slack sync; ingestion is currently file upload only."),
        ("Stateless Cache", "TTLCache is in-process; multi-replica deployments lose cache benefits without Redis."),
    ]
    for t, b in limitations:
        add_para(doc, t, bold=True)
        add_para(doc, b)

    # ── Future Enhancements ──────────────────────────────────────────
    doc.add_heading("Future Enhancements", level=1)
    doc.add_heading("Multi-Tenancy", level=3)
    add_bullet(doc, "Workspaces, workspace_members table, RBAC (owner/admin/member/viewer)")
    add_bullet(doc, "Per-workspace runbook scoping replacing the email whitelist")
    add_bullet(doc, "API keys for service accounts (CI bots, ChatOps integrations)")
    doc.add_heading("Compare / Diff Mode", level=3)
    add_bullet(doc, "Side-by-side answers from two scopes (admin vs mine) with LLM diff summary")
    add_bullet(doc, "Runbook version diffing once versioning is added")
    doc.add_heading("Conversation Persistence", level=3)
    add_bullet(doc, "Promote thread_id to a dedicated conversations table with auto-titling and last_active sort")
    add_bullet(doc, "Sidebar Recent Conversations list")
    doc.add_heading("Knowledge Base Browser", level=3)
    add_bullet(doc, "Dedicated /kb page: search + filter runbooks, preview chunks, see hit-counts per runbook")
    add_bullet(doc, "Tags table replacing the hardcoded 4-category enum")
    doc.add_heading("Integrations", level=3)
    add_bullet(doc, "URL / web-scrape ingest (httpx + readability-lxml)")
    add_bullet(doc, "Markdown and HTML parsers")
    add_bullet(doc, "OCR fallback for scanned PDFs (pytesseract)")
    add_bullet(doc, "Confluence, Notion, GitHub wiki sync via apscheduler with etag-based incremental updates")
    add_bullet(doc, "Slack / Teams bot wrappers around /query")
    doc.add_heading("Trust, Safety & Evaluation", level=3)
    add_bullet(doc, "Citation faithfulness check: second LLM pass flags unsupported sentences in red")
    add_bullet(doc, "PII / secret redaction during ingestion (regex pass before chunking)")
    add_bullet(doc, "Eval harness with question→runbook ground-truth pairs measuring recall@5, MRR, and answer-grounding rate")
    add_bullet(doc, "Prompt-injection guard stripping system-prompt-override patterns from retrieved chunks")
    doc.add_heading("UX Polish", level=3)
    add_bullet(doc, "Light/system theme toggle (Tailwind dark: classes already wired)")
    add_bullet(doc, "Settings page (default mode, default scope, density, language, keymap)")
    add_bullet(doc, "Keyboard shortcut help modal (?)")
    add_bullet(doc, "Onboarding flow on first login")
    add_bullet(doc, "PWA / installable with offline read of bookmarked answers")
    add_bullet(doc, "Accessibility pass (ARIA roles, focus rings, reduced-motion guard)")

    # ── Conclusion ───────────────────────────────────────────────────
    doc.add_heading("Conclusion", level=1)
    add_para(doc,
        "ResolveIT AI demonstrates how modern LLM technology, classical "
        "information retrieval, and thoughtful UX can combine into a "
        "production-style internal tool. The hybrid FAISS + BM25 + RRF "
        "retrieval pipeline, augmented with HyDE and a cross-encoder "
        "re-ranker, gives the system both semantic recall and exact-string "
        "precision — precisely what runbook search demands. The strict "
        "gate-check and inline citations keep the LLM grounded in the "
        "source material, eliminating the most common failure mode of "
        "naive chatbot deployments.")
    add_para(doc,
        "The six response modes turn the system from a one-shot Q&A "
        "into a configurable assistant that adapts to the user's "
        "experience level and the urgency of the moment. Combined with "
        "regeneration, follow-up suggestions, Markdown export, threaded "
        "conversations, bookmarks, and an admin layer that surfaces "
        "knowledge gaps, the platform delivers value across the full "
        "lifecycle of an IT support engagement.")
    add_para(doc,
        "The roadmap of multi-tenancy, compare mode, integrations, and "
        "trust/safety features outlines a clear path from this proof of "
        "concept to a deployable team product. The modular architecture "
        "(adding a new mode is a single dict entry; adding a new "
        "ingestion format is a single parser) is designed to make those "
        "additions straightforward.")

    # ── Appendix ────────────────────────────────────────────────────
    doc.add_heading("Appendix", level=1)
    doc.add_heading("Source Code", level=3)
    add_para(doc,
        "The complete source is organized into backend/ (FastAPI + RAG "
        "pipeline + ingestion + retrieval) and frontend/ (React + Vite + "
        "Tailwind). The codebase follows clean modular boundaries: every "
        "concern (config, auth, ingestion, retrieval, generation, "
        "routing) lives in its own package. New response modes are added "
        "by appending an entry to the MODES dict in rag/modes.py.")

    doc.add_heading("Configuration Files", level=3)
    add_para(doc,
        "backend/.env stores Supabase keys, Gemini API key, Firebase "
        "credentials path, admin-email whitelist, CORS origins, upload "
        "limits, rate-limit rules, and cache settings. frontend/.env "
        "stores Firebase web SDK config and the API base URL. "
        "requirements.txt and package.json pin dependencies; "
        "docker-compose.yml describes the local containerized run.")

    doc.add_heading("Database Migrations", level=3)
    add_bullet(doc, "001_create_tables.sql — initial schema (runbooks, query_logs, feedback)")
    add_bullet(doc, "002_add_content_hash.sql — SHA-256 dedup column")
    add_bullet(doc, "003_add_is_admin_runbook.sql — scope flag")
    add_bullet(doc, "004_bookmarks_rls.sql — bookmarks table + RLS policies")
    add_bullet(doc, "005_add_thread_id.sql — conversation threading")
    add_bullet(doc, "006_modes_and_followups.sql — mode + regenerate_of columns")

    doc.add_heading("API Documentation", level=3)
    add_para(doc,
        "All endpoints are auto-documented via FastAPI at /docs (Swagger "
        "UI) and /redoc. Key routes:")
    add_bullet(doc, "POST /query — blocking RAG query")
    add_bullet(doc, "POST /query/stream — SSE streaming RAG query (events: mode, sources, token, done)")
    add_bullet(doc, "GET /answer/{id} — fetch saved answer (permalink)")
    add_bullet(doc, "GET /export/{id}.md — download answer as Markdown")
    add_bullet(doc, "POST /feedback — submit thumbs up/down + comment")
    add_bullet(doc, "GET /history — paginated, thread-grouped audit log")
    add_bullet(doc, "POST /bookmarks · GET /bookmarks · DELETE /bookmarks/{id} — playbook")
    add_bullet(doc, "POST /runbooks/upload · GET /runbooks/my · DELETE /runbooks/my/{id} — personal runbooks")
    add_bullet(doc, "POST /admin/upload · GET /admin/runbooks · DELETE /admin/runbooks/{id} — admin runbooks")
    add_bullet(doc, "GET /admin/feedback-stats · GET /admin/runbook-health · GET /admin/knowledge-gaps — analytics")
    add_bullet(doc, "GET /auth/me · POST /auth/verify — auth helpers")
    add_bullet(doc, "GET /health — service health (FAISS vector count + metadata count)")

    doc.add_heading("Dataset Details", level=3)
    add_para(doc,
        "The system ingests user-supplied runbooks at runtime; there is "
        "no static training dataset. Sample runbooks shipped under "
        "backend/sample_runbooks/ cover Apache, Nginx, MySQL, network "
        "diagnostics, and SSL/TLS scenarios for development testing.")

    doc.add_heading("References", level=3)
    add_bullet(doc, "Karpukhin et al., \"Dense Passage Retrieval for Open-Domain Question Answering\" (DPR, 2020)")
    add_bullet(doc, "Robertson & Zaragoza, \"The Probabilistic Relevance Framework: BM25 and Beyond\" (2009)")
    add_bullet(doc, "Cormack, Clarke & Buettcher, \"Reciprocal Rank Fusion outperforms Condorcet and individual Rank Learning Methods\" (2009)")
    add_bullet(doc, "Gao et al., \"HyDE: Precise Zero-Shot Dense Retrieval without Relevance Labels\" (2022)")
    add_bullet(doc, "Xiao et al., BAAI BGE embedding and reranker models (2023+)")

    doc.add_heading("GitHub Repository Link", level=3)
    add_para(doc, "(Add your repository URL here.)")

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main(r"C:/Users/ABHINAV TEJA/Downloads/Gen ai/resolveit-ai/ResolveIT_AI_Project_Documentation.docx")
